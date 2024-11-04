"Create DOCX file."

from icecream import ic

import copy
import datetime
import io
import os

import docx
import docx.oxml

import constants
import utils

Tx = utils.Tx


class Creator:
    "DOCX creator."

    def __init__(self, book, references, item=None):
        self.book = book
        self.references = references
        self.item = item
        self.title = book.title
        self.subtitle = book.subtitle
        self.authors = book.authors
        self.language = book.language
        settings = book.frontmatter["docx"]
        self.title_page_metadata = settings["title_page_metadata"]
        self.page_break_level = settings["page_break_level"]
        self.footnotes_location = settings["footnotes_location"]
        self.indexed_font = settings.get("indexed_font")
        self.reference_font = settings.get("reference_font")

    def create(self):
        "Create the DOCX document; return a BytesIO instance containing it."
        # Key: fulltitle; value: dict(label, ast_children)
        self.footnotes = {}
        # References, key: refid; value: reference
        self.referenced = set()
        # Key: canonical; value: dict(id, fulltitle, ordinal)
        self.indexed = {}
        self.indexed_count = 0

        self.document = docx.Document()

        # Set the default document-wide language.
        # From https://stackoverflow.com/questions/36967416/how-can-i-set-the-language-in-text-with-python-docx
        if self.language:
            styles_element = self.document.styles.element
            rpr_default = styles_element.xpath("./w:docDefaults/w:rPrDefault/w:rPr")[0]
            lang_default = rpr_default.xpath("w:lang")[0]
            lang_default.set(docx.oxml.shared.qn("w:val"), self.language)

        # Set to A4 page size.
        section = self.document.sections[0]
        section.page_height = docx.shared.Mm(297)
        section.page_width = docx.shared.Mm(210)
        section.left_margin = docx.shared.Mm(25.4)
        section.right_margin = docx.shared.Mm(25.4)
        section.top_margin = docx.shared.Mm(25.4)
        section.bottom_margin = docx.shared.Mm(25.4)
        section.header_distance = docx.shared.Mm(12.7)
        section.footer_distance = docx.shared.Mm(12.7)

        # Create style for code.
        style = self.document.styles.add_style(
            constants.CODE_STYLE, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH
        )
        style.base_style = self.document.styles["macro"]
        style.paragraph_format.left_indent = docx.shared.Pt(constants.CODE_LEFT_INDENT)
        style.font.name = constants.CODE_FONT

        # Create style for quote.
        style = self.document.styles.add_style(
            constants.QUOTE_STYLE, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH
        )
        style.paragraph_format.left_indent = docx.shared.Pt(constants.QUOTE_LEFT_INDENT)
        style.paragraph_format.right_indent = docx.shared.Pt(
            constants.QUOTE_RIGHT_INDENT
        )
        style.font.name = constants.QUOTE_FONT

        # Set Dublin core metadata.
        if self.language:
            self.document.core_properties.language = self.language
            self.document.core_properties.modified = datetime.datetime.now()
        # XXX authors

        self.current_text = None
        self.footnote_paragraph = None

        if self.item is None:
            self.write_title_page()
            self.write_toc()
            items = list(self.book.items)
        elif self.item.is_section:
            paragraph = self.document.add_paragraph(style="Title")
            run = paragraph.add_run(self.item.title)
            run.font.size = docx.shared.Pt(24)
            run.font.bold = True
            items = list(self.item.all_items)
        elif self.item.is_text:
            paragraph = self.document.add_paragraph(style="Title")
            run = paragraph.add_run(self.item.title)
            run.font.size = docx.shared.Pt(20)
            run.font.bold = True
            items = [self.item]
        self.write_page_number()
        for item in items:
            if item.is_section:
                self.write_section(item, level=item.level)
            else:
                self.write_text(item, level=item.level)
            self.write_footnotes_chapter(item)
        self.write_footnotes_book()
        self.write_references()
        self.write_indexed()

        output = io.BytesIO()
        self.document.save(output)
        return output

    def write_title_page(self):
        paragraph = self.document.add_paragraph(style="Title")
        run = paragraph.add_run(self.title)
        run.font.size = docx.shared.Pt(28)
        run.font.bold = True

        if self.subtitle:
            paragraph = self.document.add_paragraph(style="Heading 1")
            paragraph.paragraph_format.space_after = docx.shared.Pt(20)
            paragraph.add_run(self.subtitle)

        paragraph = self.document.add_paragraph(style="Heading 2")
        paragraph.paragraph_format.space_after = docx.shared.Pt(10)
        for author in self.authors:
            paragraph.add_run(author)
            if author != self.authors[-1]:
                paragraph.add_run(", ")

        self.render(self.book.ast, initialize=True)

        if self.title_page_metadata:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.space_before = docx.shared.Pt(50)
            status = str(
                min([t.status for t in self.book.all_texts] + [max(constants.STATUSES)])
            )
            paragraph.add_run(f"{Tx('Status')}: {Tx(status)}")
            now = datetime.datetime.now().strftime(constants.DATETIME_ISO_FORMAT)
            self.document.add_paragraph(f"{Tx('Created')}: {now}")

    def write_toc(self):
        "Write table of contents."
        # XXX
        pass

    def write_page_number(self):
        "Display page number in the header."
        # From https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
        paragraph = self.document.sections[-1].header.paragraphs[0]
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        fldChar1 = docx.oxml.OxmlElement("w:fldChar")
        fldChar1.set(docx.oxml.ns.qn("w:fldCharType"), "begin")
        instrText = docx.oxml.OxmlElement("w:instrText")
        instrText.set(docx.oxml.ns.qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = docx.oxml.OxmlElement("w:fldChar")
        fldChar2.set(docx.oxml.ns.qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def write_section(self, section, level):
        if level <= self.page_break_level:
            self.document.add_page_break()
        self.write_heading(section.heading, level)
        for item in section.items:
            if item.is_section:
                self.write_section(item, level=level + 1)
            else:
                self.write_text(item, level=level + 1)

    def write_text(self, text, level):
        if level <= self.page_break_level:
            self.document.add_page_break()
        if not text.frontmatter.get("suppress_title"):
            self.write_heading(text.heading, level)
        self.current_text = text
        self.render(text.ast, initialize=True)
        self.write_footnotes_text(text)

    def write_heading(self, heading, level):
        level = min(level, constants.MAX_H_LEVEL)
        h = constants.H_LOOKUP[level]
        paragraph = self.document.add_paragraph(style=f"Heading {level}")
        paragraph.paragraph_format.left_indent = docx.shared.Pt(h["left_margin"])
        paragraph.paragraph_format.space_after = docx.shared.Pt(h["spacing"])
        run = paragraph.add_run(heading)
        run.font.size = docx.shared.Pt(h["font"][1])

    def write_footnotes_text(self, text):
        "Footnotes at end of the text."
        if self.footnotes_location != constants.FOOTNOTES_EACH_TEXT:
            return
        try:
            footnotes = self.footnotes[text.fulltitle]
        except KeyError:
            return
        paragraph = self.document.add_heading(Tx("Footnotes"), 6)
        paragraph.paragraph_format.space_before = docx.shared.Pt(25)
        paragraph.paragraph_format.space_after = docx.shared.Pt(10)
        for entry in sorted(footnotes.values(), key=lambda e: e["number"]):
            self.footnote_paragraph = self.document.add_paragraph()
            run = self.footnote_paragraph.add_run(f"{entry['number']}. ")
            run.italic = True
            for child in entry["ast_children"]:
                self.render(child)
            self.footnote_paragraph = None

    def write_footnotes_chapter(self, item):
        "Footnote definitions at the end of a chapter."
        if self.footnotes_location != constants.FOOTNOTES_EACH_CHAPTER:
            return
        try:
            footnotes = self.footnotes[item.chapter.fulltitle]
        except KeyError:
            return
        self.document.add_page_break()
        self.write_heading(Tx("Footnotes"), 4)
        for entry in sorted(footnotes.values(), key=lambda e: e["number"]):
            self.footnote_paragraph = self.document.add_paragraph()
            run = self.footnote_paragraph.add_run(f"{entry['number']}. ")
            run.italic = True
            for child in entry["ast_children"]:
                self.render(child)
            self.footnote_paragraph = None

    def write_footnotes_book(self):
        "Footnote definitions as a separate section at the end of the book."
        if self.footnotes_location != constants.FOOTNOTES_END_OF_BOOK:
            return
        self.document.add_page_break()
        self.write_heading(Tx("Footnotes"), 1)
        for item in self.book.items:
            footnotes = self.footnotes.get(item.fulltitle, {})
            if not footnotes:
                continue
            self.write_heading(item.heading, 2)
            for entry in sorted(footnotes.values(), key=lambda e: e["number"]):
                self.footnote_paragraph = self.document.add_paragraph()
                run = self.footnote_paragraph.add_run(f"{entry['number']}. ")
                run.italic = True
                for child in entry["ast_children"]:
                    self.render(child)
                self.footnote_paragraph = None

    def write_references(self):
        if not self.referenced:
            return
        self.document.add_page_break()
        self.write_heading(Tx("References"), 1)
        for refid in sorted(self.referenced):
            try:
                reference = self.references[refid]
            except KeyError:
                continue
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(reference["name"])
            run.bold = True
            paragraph.add_run("  ")
            self.write_reference_authors(paragraph, reference)
            try:
                method = getattr(self, f"write_reference_{reference['type']}")
            except AttributeError:
                print("unknown", reference["type"])
            else:
                method(paragraph, reference)
            self.write_reference_external_links(paragraph, reference)

    def write_reference_authors(self, paragraph, reference):
        count = len(reference["authors"])
        for pos, author in enumerate(reference["authors"]):
            if pos > 0:
                if pos == count - 1:
                    paragraph.add_run(" & ")
                else:
                    paragraph.add_run(", ")
            paragraph.add_run(utils.short_name(author))

    def write_reference_article(self, paragraph, reference):
        paragraph.add_run(f"({reference['year']}) ")
        paragraph.add_run(utils.full_title(reference))
        try:
            run = paragraph.add_run(f"{reference['journal']} ")
            run.font.italic = True
        except KeyError:
            pass
        try:
            paragraph.add_run(f"{reference['volume']} ")
        except KeyError:
            pass
        else:
            try:
                paragraph.add_run(f"({reference['number']})")
            except KeyError:
                pass
        try:
            paragraph.add_run(f": pp. {reference['pages'].replace('--', '-')}.")
        except KeyError:
            pass

    def write_reference_book(self, paragraph, reference):
        paragraph.add_run(f"({reference['year']}). ")
        run = paragraph.add_run(utils.full_title(reference))
        run.font.italic = True
        try:
            paragraph.add_run(f"{reference['publisher']}. ")
        except KeyError:
            pass

    def write_reference_link(self, paragraph, reference):
        paragraph.add_run(f"({reference['year']}). ")
        title = utils.full_title(reference)
        paragraph.add_run(title)
        try:
            add_hyperlink(paragraph, reference["url"], title)
        except KeyError:
            pass
        try:
            paragraph.add_run(f" Accessed {reference['accessed']}.")
        except KeyError:
            pass

    def write_reference_external_links(self, paragraph, reference):
        any_item = False
        if reference.get("url"):
            add_hyperlink(paragraph, reference["url"], reference["url"])
            any_item = True
        for key, (label, template) in constants.REFERENCE_LINKS.items():
            try:
                value = reference[key]
                if any_item:
                    paragraph.add_run(", ")
                else:
                    paragraph.add_run("  ")
                add_hyperlink(
                    paragraph, template.format(value=value), f"{label}:{value}"
                )
                any_item = True
            except KeyError:
                pass

    def write_indexed(self):
        if not self.indexed:
            return
        self.document.add_page_break()
        self.write_heading(Tx("Index"), 1)
        items = sorted(self.indexed.items(), key=lambda i: i[0].lower())
        for canonical, entries in items:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(canonical)
            run.bold = True
            paragraph.add_run("  ")
            entries.sort(key=lambda e: e["ordinal"])
            for entry in entries:
                paragraph.add_run(entry["heading"])
                if entry is not entries[-1]:
                    paragraph.add_run(", ")

    def render(self, ast, initialize=False):
        if initialize:
            self.list_stack = []
            self.style_stack = ["Normal"]
            self.bold = False
            self.italic = False
            self.subscript = False
            self.superscript = False
        try:
            method = getattr(self, f"render_{ast['element']}")
        except AttributeError:
            print("Could not handle ast", ast)
        else:
            method(ast)

    def render_document(self, ast):
        self.prev_blank_line = False
        for child in ast["children"]:
            self.render(child)

    def render_heading(self, ast):
        # XXX Limited implementation; just handles one child of raw text.
        text = ast["children"][0]["children"]
        self.write_heading(text, ast["level"])

    def render_paragraph(self, ast):
        if self.footnote_paragraph:
            self.paragraph = self.footnote_paragraph
        else:
            self.paragraph = self.document.add_paragraph()
        if self.list_stack:
            data = self.list_stack[-1]
            depth = min(3, data["depth"])  # Max list depth in predef styles.
            if data["first_paragraph"]:
                if data["ordered"]:
                    if depth == 1:
                        style = self.document.styles["List Number"]
                    else:
                        style = self.document.styles[f"List Number {depth}"]
                else:
                    if depth == 1:
                        style = self.document.styles["List Bullet"]
                    else:
                        style = self.document.styles[f"List Bullet {depth}"]
            else:
                if depth == 1:
                    style = self.document.styles["List Continue"]
                else:
                    style = self.document.styles[f"List Continue {depth}"]
            data["first_paragraph"] = False
            self.paragraph.style = style
        else:
            self.paragraph.style = self.style_stack[-1]
        for child in ast["children"]:
            self.render(child)

    def render_raw_text(self, ast):
        line = ast["children"]
        line = line.rstrip("\n")
        run = self.paragraph.add_run(line)
        if self.bold:
            run.font.bold = True
        if self.italic:
            run.font.italic = True
        if self.subscript:
            run.font.subscript = True
        if self.superscript:
            run.font.superscript = True

    def render_blank_line(self, ast):
        pass

    def render_quote(self, ast):
        self.style_stack.append(constants.QUOTE_STYLE)
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_code_span(self, ast):
        run = self.paragraph.add_run(ast["children"])
        run.style = self.document.styles["Macro Text Char"]

    def render_code_block(self, ast):
        self.paragraph = self.document.add_paragraph(style=constants.CODE_STYLE)
        self.style_stack.append(constants.CODE_STYLE)
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_fenced_code(self, ast):
        self.paragraph = self.document.add_paragraph(style=constants.CODE_STYLE)
        self.style_stack.append(constants.CODE_STYLE)
        for child in ast["children"]:
            self.render(child)
        self.style_stack.pop()

    def render_emphasis(self, ast):
        self.italic = True
        for child in ast["children"]:
            self.render(child)
        self.italic = False

    def render_strong_emphasis(self, ast):
        self.bold = True
        for child in ast["children"]:
            self.render(child)
        self.bold = False

    def render_subscript(self, ast):
        self.subscript = True
        for child in ast["children"]:
            self.render(child)
        self.subscript = False

    def render_superscript(self, ast):
        self.superscript = True
        for child in ast["children"]:
            self.render(child)
        self.superscript = False

    def render_emdash(self, ast):
        self.paragraph.add_run(constants.EM_DASH)

    def render_line_break(self, ast):
        if ast.get("soft"):
            self.paragraph.add_run(" ")
        else:
            self.paragraph.add_run("\n")

    def render_thematic_break(self, ast):
        paragraph = self.document.add_paragraph(constants.EM_DASH * 20)
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def render_link(self, ast):
        # This handles only raw text within a link, nothing else.
        raw_text = []
        for child in ast["children"]:
            if child["element"] == "raw_text":
                raw_text.append(child["children"])
        add_hyperlink(self.paragraph, ast["dest"], "".join(raw_text))

    def render_list(self, ast):
        data = dict(
            ordered=ast["ordered"],
            bullet=ast["bullet"],  # Currently useless.
            start=ast["start"],  # Currently useless.
            tight=ast["tight"],  # Currently useless.
            count=0,  # Currently useless.
            depth=len(self.list_stack) + 1,
        )
        self.list_stack.append(data)
        for child in ast["children"]:
            self.render(child)
        self.list_stack.pop()

    def render_list_item(self, ast):
        data = self.list_stack[-1]
        data["count"] += 1  # Currently useless.
        data["first_paragraph"] = True
        for child in ast["children"]:
            self.render(child)

    def render_indexed(self, ast):
        entries = self.indexed.setdefault(ast["canonical"], [])
        self.indexed_count += 1
        entries.append(
            dict(
                id=f"i{self.indexed_count}",
                ordinal=self.current_text.ordinal,
                fulltitle=self.current_text.fulltitle,
                heading=self.current_text.heading,
            )
        )
        run = self.paragraph.add_run(ast["term"])
        if self.indexed_font == constants.ITALIC:
            run.italic = True
        elif self.indexed_font == constants.BOLD:
            run.bold = True
        elif self.indexed_font == constants.UNDERLINE:
            run.underline = True

    def render_footnote_ref(self, ast):
        # The label is used only for lookup; number is used for output.
        label = ast["label"]
        if self.footnotes_location == constants.FOOTNOTES_EACH_TEXT:
            entries = self.footnotes.setdefault(self.current_text.fulltitle, {})
            number = len(entries) + 1
            key = label
        elif self.footnotes_location in (
            constants.FOOTNOTES_EACH_CHAPTER,
            constants.FOOTNOTES_END_OF_BOOK,
        ):
            fulltitle = self.current_text.chapter.fulltitle
            entries = self.footnotes.setdefault(fulltitle, {})
            number = len(entries) + 1
            key = f"{fulltitle}-{label}"
        entries[key] = dict(label=label, number=number)
        run = self.paragraph.add_run(str(number))
        run.font.superscript = True
        run.font.bold = True

    def render_footnote_def(self, ast):
        label = ast["label"]
        if self.footnotes_location == constants.FOOTNOTES_EACH_TEXT:
            fulltitle = self.current_text.fulltitle
            key = label
        elif self.footnotes_location in (
            constants.FOOTNOTES_EACH_CHAPTER,
            constants.FOOTNOTES_END_OF_BOOK,
        ):
            fulltitle = self.current_text.chapter.fulltitle
            key = f"{fulltitle}-{label}"
        self.footnotes[fulltitle][key]["ast_children"] = ast["children"]

    def render_reference(self, ast):
        self.referenced.add(ast["id"])
        run = self.paragraph.add_run(ast["name"])
        if self.reference_font == constants.ITALIC:
            run.italic = True
        elif self.reference_font == constants.BOLD:
            run.bold = True
        elif self.reference_font == constants.UNDERLINE:
            run.underline = True


# From https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410
def add_hyperlink(paragraph, url, text, color="2222FF", underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value.
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    # Create the w:hyperlink tag and add needed values.
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(
        docx.oxml.shared.qn("r:id"),
        r_id,
    )

    # Create a w:r element.
    new_run = docx.oxml.shared.OxmlElement("w:r")

    # Create a new w:rPr element.
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Add color if it is given.
    if not color is None:
        c = docx.oxml.shared.OxmlElement("w:color")
        c.set(docx.oxml.shared.qn("w:val"), color)
        rPr.append(c)

    # Remove underlining if it is requested.
    # XXX Does not seem to work? /Per Kraulis
    if not underline:
        u = docx.oxml.shared.OxmlElement("w:u")
        u.set(docx.oxml.shared.qn("w:val"), "none")
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element.
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink
